import os
from datetime import datetime
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH 

OUTPUT_DIR = "output_documents"

def build_word_document(title: str, content: dict[str, str]) -> str:
    os.makedirs(OUTPUT_DIR, exist_ok=True)

    doc = Document()

    heading = doc.add_heading(title, level=0)                                                                         
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(f"Generated on: {datetime.now().strftime('%B %d, %Y %H:%M')}")                                  
    doc.add_paragraph("")

    for section_title, section_content in content.items():                                                            
        doc.add_heading(section_title, level=1)
        para = doc.add_paragraph(section_content)                                                                     
        para.paragraph_format.space_after = Pt(12) 

    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    filename = f"document_{timestamp}.docx"
    file_path = os.path.join(OUTPUT_DIR, filename)                                                                    
    doc.save(file_path)                     
                                                                                                                    
    return file_path
